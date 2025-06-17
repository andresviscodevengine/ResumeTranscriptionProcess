from typing import List, Dict, Optional
import fitz
import functions_framework
import os
from google.cloud import storage
import vertexai
from vertexai.language_models import TextEmbeddingModel, TextGenerationModel
from flask import Request, Flask, request
import json
from vertexai.generative_models import GenerativeModel
from google.cloud import aiplatform
from docx import Document
import re
from io import BytesIO
import numpy as np


#### BLOQUE DEBUG LOCAL #####
# app = Flask(__name__)


TEMPERATURE = 0.5
TOP_K = 40
TOP_P = 0.90

JOB_INTERVIEW_KEYWORDS = [
        "interview", "position", "experience", "resume", "skills", "cv",
        "salary", "company", "role", "aspirations", "availability",
        "expectations", "challenges", "team", "projects", "future",
        "questions", "thank you for your time", "opportunity", "candidate",
        "onboarding", "background", "strengths", "weaknesses", "tell me about yourself"
    ]
MIN_KEYWORD_MATCHES = 1


# app = Flask(__name__)
# Configuraciones
PROJECT_ID = "celtic-tendril-455220-v1"
BUCKET_NAME = "transcription_poc_uploads_raw"
BUCKET_DESTINO = "transcription_poc_processed"
EMBEDDING_MODEL = "text-embedding-large-exp-03-07"
LLM_MODEL = "gemini-1.5-flash-002"
REGION = "us-central1"
PROMPT = "prompt.prompt"
# Inicializar Vertex AI


def limpiar_transcripcion_texto(pdf_bytes: bytes, candidate_name: str) -> str:
    """
    Cleans and processes the text of a transcription in PDF format.

    Args:
        pdf_bytes (bytes): Content of the PDF file in bytes.
        candidate_name (str): Candidate's name to filter relevant lines.

    Returns:
        str: Cleaned and processed text.
    """
    try:
        # Leer el PDF y extraer el texto completo
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        full_text = "\n".join(page.get_text() for page in pdf_doc)
    except Exception as e:
        print(f"Error al leer PDF: {e}")
        full_text = pdf_bytes.decode("utf-8", errors="ignore")

    # El resto sigue igual que en tu versión original:
    full_text = re.sub(r"\b\d{2}:\d{2}:\d{2}\b", "", full_text)
    full_text = re.sub(r"\n\s*\n", "\n", full_text)

    candidate_name_lower = candidate_name.lower()
    full_text = re.sub(candidate_name, candidate_name_lower, full_text, flags=re.IGNORECASE)

    fillers = [
        r"\buh\b", r"\bum\b", r"\byou know\b", r"\bokay\b", r"\bsorry\b", r"\bwell\b",
        r"\bso\b", r"\bi mean\b", r"\blet me see\b", r"\bjust\b", r"\blike\b", r"\bkind of\b", r"\bperfect\b", r"\byeah\b",
    ]
    full_text = re.sub("|".join(fillers), "", full_text, flags=re.IGNORECASE)
    full_text = re.sub(r"\s{2,}", " ", full_text)

    lines = full_text.splitlines()
    filtered_lines = [
        line.strip()
        for line in lines
        if line.strip().lower().startswith(candidate_name_lower + ":")
    ]

    return "\n".join(filtered_lines)

def extract_reclut(file_name: str) -> str:
    """
    Extracts the recruiter/codename that comes after the last '-' and before '.pdf'.
    """
    file_name = file_name.lower()
    
    final_match = re.search(r'-([^-]+)\.pdf$', file_name)
    if final_match:
        reclut = final_match.group(1).strip()
        return reclut
    else:
        raise ValueError("No se pudo extraer el reclutador del nombre del archivo.")

def extract_candidate_name(file_name: str) -> str:
    """
    Extrae el nombre del candidato desde el nombre del archivo.
    Si no hay paréntesis, retorna 'candidato_sin_nombre'.
    """
    # Validar formato esperado
    
    # if not file_name.lower().endswith(".pdf"):
    #     raise ValueError("Formato de archivo inesperado")

    # Buscar la parte final: '-[nombre].pdf'
    print(str(file_name))
    # final_match = re.search(r'-([^-]+)\.pdf$', file_name)
    final_match = re.search(r'\(([^)]+)\)', file_name)

    if not final_match:
        return "nonamecandidate"    
    else:
        nombre = final_match.group(1)
        return nombre
    # cutoff_index = final_match.start()
    # prefix = file_name[:cutoff_index]

    # # Buscar el último grupo entre paréntesis antes del sufijo
    # parens_matches = re.findall(r"\(([^)]+)\)", prefix)
    # if parens_matches:
    #     return parens_matches[-1].strip()
    
    # No hay nombre entre paréntesis → nombre no disponible
    

def download_and_extract_text(bucket_name: str, file_name: str) -> Optional[str]:
    """
    Downloads a file from GCS and extracts its text content.
    It handles both PDF and plain text files.
    """
    try:
        print(f"Attempting to download and process '{file_name}' from bucket '{bucket_name}'...")
        client = storage.Client()
        bucket = client.get_bucket(bucket_name)
        blob = bucket.blob(file_name)

        file_bytes = blob.download_as_bytes()
        print(f"File '{file_name}' downloaded successfully ({len(file_bytes)} bytes).")

        if file_name.lower().endswith('.pdf'):
            print("PDF file detected. Extracting text with PyMuPDF...")
            text_content = ""
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf_doc:
                for page_num, page in enumerate(pdf_doc):
                    text_content += page.get_text()
            print("Text extracted from PDF successfully.")
            return text_content
        else:
            print("Text file detected. Decoding as UTF-8.")
            return file_bytes.decode('utf-8')

    except Exception as e:
        print(f"Error during file download or text extraction for '{file_name}': {e}")
        return None

def clean_and_extract_dialogue_segment(full_text_content: str) -> str:
    """
    Cleans the text content to extract only the dialogue portion.
    It finds the keyword "Transcript" (case-insensitive, with optional variations) and returns all text that follows the last such marker before the first speaker label.
    """
    if not full_text_content:
        print("Warning: Full text content is empty. Cannot extract dialogue segment.")
        return ""

    # Regex for variations of "Transcript"
    transcript_pattern = re.compile(r"(?:\)[ ]*)?- Transcript|-\s*Transcript|Transcript", re.IGNORECASE)
    # Patron: línea que comienza con un nombre (solo letras y espacios, sin números ni caracteres raros), seguido exactamente de ": "
    speaker_pattern = re.compile(r"^([A-Za-zÁÉÍÓÚáéíóúÑñüÜ\s]+):\s+[A-Za-zÁÉÍÓÚáéíóúÑñüÜ]", re.MULTILINE)

    # Find all "Transcript" markers
    transcript_matches = list(transcript_pattern.finditer(full_text_content))
    if not transcript_matches:
        print("Warning: 'Transcript' keyword not found in the document. Dialogue segment will be empty.")
        return ""

    # Find the first speaker label after all transcript markers
    speaker_match = speaker_pattern.search(full_text_content)
    if not speaker_match:
        print("Warning: No speaker label found after 'Transcript'. Dialogue segment will be empty.")
        return ""

    # Find the last transcript marker before the first speaker label
    last_valid_marker = None
    for m in transcript_matches:
        if m.end() <= speaker_match.start():
            last_valid_marker = m
        else:
            break

    if last_valid_marker:
        dialogue_start_index = last_valid_marker.end()
        print(f"Found 'Transcript' marker at index {last_valid_marker.start()}. Extracting subsequent text.")
        return full_text_content[dialogue_start_index:].strip()
    else:
        # If all markers are after the first speaker, fallback to the first marker
        dialogue_start_index = transcript_matches[0].end()
        print(f"Using first 'Transcript' marker at index {transcript_matches[0].start()} (no marker before first speaker).")
        return full_text_content[dialogue_start_index:].strip()

def is_job_interview(text: str, keywords: List[str], min_keyword_matches: int) -> bool:
    """
    Determines if the text corresponds to a job interview based on keywords.
    """
    if not text:
        print("Warning: Text for keyword check is empty.")
        return False
    text_lower = text.lower()
    found_keywords_set = set()
    for keyword in keywords:
        if keyword.lower() in text_lower:
            found_keywords_set.add(keyword.lower())
    num_found = len(found_keywords_set)
    print(f"Keyword check: Found {num_found} distinct keywords out of {min_keyword_matches} required.")
    if num_found >= min_keyword_matches:
        print(f"Matching keywords: {list(found_keywords_set)}")
        return True
    return False


def comparar_cadenas_por_palabras(cadena_entrada: str, cadena_referencia: str, umbral_mayoria: float = 0.5) -> bool:
    """
    Compara dos cadenas de texto basándose en las palabras que contienen.

    La función devuelve True si se cumple alguna de las siguientes condiciones:
    1. Ambas cadenas, después de normalizar (minúsculas, palabras únicas), contienen
       exactamente el mismo conjunto de palabras.
    2. La mayoría de las palabras únicas en la `cadena_entrada` (después de normalizar)
       también están presentes en la `cadena_referencia` (después de normalizar).
       "Mayoría" se define por defecto como más del 50% (umbral > 0.5).

    En caso contrario, o si alguna cadena está vacía y la otra no, devuelve False.
    Si ambas cadenas están vacías (o solo contienen espacios), devuelve True.

    Args:
        cadena_entrada (str): La cadena de texto a evaluar.
        cadena_referencia (str): La cadena de texto de referencia para la comparación.
        umbral_mayoria (float, optional): El umbral para considerar una "mayoría" de coincidencias.
                                         Debe estar entre 0.0 y 1.0. Por defecto es 0.5,
                                         lo que significa que más del 50% de las palabras deben coincidir.

    Returns:
        bool: True si las condiciones de coincidencia se cumplen, False en caso contrario.
    """

    # Paso 1: Preprocesamiento y normalización
    # Convertir a minúsculas, dividir en palabras y obtener conjuntos de palabras únicas.
    # split() sin argumentos maneja múltiples espacios y espacios al inicio/final.
    palabras_entrada = set(cadena_entrada.lower().split())
    palabras_referencia = set(cadena_referencia.lower().split())

    # Manejar casos de cadenas completamente vacías o solo con espacios
    if not palabras_entrada and not palabras_referencia:
        # Ambas cadenas están vacías (o solo contenían espacios)
        return True
    if not palabras_entrada or not palabras_referencia:
        # Una cadena está vacía (o solo espacios) y la otra no.
        # No pueden coincidir ni tener una mayoría de coincidencias.
        return False

    # Condición 1: Coincidencia exacta de los conjuntos de palabras.
    # Esto implica que tienen las mismas palabras únicas y la misma cantidad de palabras únicas.
    # Cubre el caso "si coincide todas pero distinto orden, está ok" y "misma cantidad de palabras".
    if palabras_entrada == palabras_referencia:
        return True

    # Condición 2: Verificar si la mayoría de las palabras de `cadena_entrada` están en `cadena_referencia`.
    # Calcular la cantidad de palabras comunes.
    palabras_comunes = palabras_entrada.intersection(palabras_referencia)
    num_palabras_comunes = len(palabras_comunes)
    num_palabras_entrada_unicas = len(palabras_entrada)

    # Si no hay palabras únicas en la entrada (ya cubierto arriba, pero por seguridad),
    # no se puede calcular el ratio.
    if num_palabras_entrada_unicas == 0:
        return False # Ya debería haber sido capturado por el "not palabras_entrada" anterior.

    ratio_coincidencia = num_palabras_comunes / num_palabras_entrada_unicas

    # "si no coincide la mayoria, devuelve false"
    # "o ninguna conicide" (esto resultaría en ratio_coincidencia = 0)
    if ratio_coincidencia > umbral_mayoria:
        return True
    
    return False





def parse_interview_dialogue(dialogue_text: str, primary_speaker_known_name: str) -> List[Dict[str, str]]:
    """
    Parses the dialogue text and extracts conversation turns, using actual speaker names.
    The primary_speaker_known_name is one of the main speakers (e.g., interviewer).
    Other speaker names are derived from the transcript.

    Args:
        dialogue_text (str): The clean text containing only the dialogue.
        primary_speaker_known_name (str): The known name of one of the primary speakers.
                                         This name will be used as is in the output for this speaker.

    Returns:
        List[Dict[str, str]]: A list of dictionaries, each representing a turn.
                                 e.g., [{'speaker': 'Actual Name', 'text': '...'}]
    """
    if not dialogue_text:
        print("Warning: Dialogue text for parsing is empty.")
        return []

    parsed_dialogue = []
    primary_speaker_known_name_lower = primary_speaker_known_name.lower().strip()
    
    # To store the name of the other main speaker (e.g., interviewee) once identified.
    # We store the first encountered name and its casing for consistency.
    other_speaker_identified_name = None 

    # Regex to split by speaker labels (e.g., "Name:", "[timestamp] Name:")
    # This pattern captures the speaker label itself.
    speaker_pattern = re.compile(
        r'((?:\[\d{2}:\d{2}:\d{2}(?:\.\d+)?\]\s*)?[\w\s\-\.]+\s*:\s*)',
        re.IGNORECASE
    )
    
    # Split the dialogue by speaker labels. `parts` will contain text before the first label,
    # then label1, then text1, then label2, then text2, and so on.
    parts = speaker_pattern.split(dialogue_text)
    
    current_speaker_label_from_parts = None # Stores the full speaker label like "massucatto jean:"
    
    # Iterate through the parts. Start from index 1 because parts[0] is any text
    # before the first speaker label (or empty if dialogue starts with a label).
    for i in range(1, len(parts)):
        part_content = parts[i].strip()
        if not part_content: # Skip empty parts that might result from the split
            continue

        # Check if the current part is a speaker label based on the regex
        is_speaker_label = bool(speaker_pattern.fullmatch(part_content))

        if is_speaker_label:
            current_speaker_label_from_parts = part_content
        elif current_speaker_label_from_parts: 
            # This part is dialogue text belonging to the previously identified current_speaker_label_from_parts
            dialogue_segment = part_content
            
            # Extract the actual name from the label (stripping timestamps).
            # e.g., from "massucatto jean:" or "[00:00:00] massucatto jean:"
            # Group 1 of this match will be the name part.
            speaker_name_match = re.match(
                r'(?:\[.*?\]\s*)?([\w\s\-\.]+)\s*:', # Capture group 1 is the name
                current_speaker_label_from_parts,
                re.IGNORECASE
            )
            if speaker_name_match:
                detected_speaker_actual_name = speaker_name_match.group(1).strip()
                speaker_display_name = "" # This will be the name used in the output
            
                # Check if the detected speaker is the primary known speaker
                speaker_name_match_ = comparar_cadenas_por_palabras(
                current_speaker_label_from_parts, primary_speaker_known_name)
                if speaker_name_match_:#detected_speaker_actual_name.lower() == primary_speaker_known_name_lower:
                    speaker_display_name = primary_speaker_known_name # Use the exact casing provided for the primary speaker
                else:
                    # This is another speaker
                    if other_speaker_identified_name is None:
                        # First time we encounter this "other" speaker. Store their name as identified.
                        other_speaker_identified_name = detected_speaker_actual_name
                        speaker_display_name = detected_speaker_actual_name # Use the name as it appeared in the transcript
                        print(f"Identified other primary speaker as: {other_speaker_identified_name}")
                    elif detected_speaker_actual_name.lower() == other_speaker_identified_name.lower():
                        # This is the same "other" speaker identified earlier. Use the consistently stored name.
                        speaker_display_name = other_speaker_identified_name
                    else:
                        # This is a new speaker name, different from the primary and the first "other" speaker.
                        # This could be a third participant or a variation. We'll label them with their actual detected name.
                        print(f"Note: New speaker detected '{detected_speaker_actual_name}', different from primary ('{primary_speaker_known_name}') and first other ('{other_speaker_identified_name}').")
                        speaker_display_name = detected_speaker_actual_name 

                if speaker_display_name and dialogue_segment: # Ensure we have a name and text
                    parsed_dialogue.append({"speaker": speaker_display_name, "text": dialogue_segment})
            
            # Reset current_speaker_label_from_parts as its text has been processed.
            # The next speaker label will set this again.
            current_speaker_label_from_parts = None 

    if not parsed_dialogue and dialogue_text:
         print("Warning: Dialogue parsing did not yield any structured turns, though dialogue text was present.")
    else:
        print(f"Dialogue parsed into {len(parsed_dialogue)} turns.")
    return parsed_dialogue

def process_transcript_file(
    bucket_name: str,
    file_name: str,
    person_a_name: str, # This will be passed as primary_speaker_known_name
    interview_keywords: List[str],
    min_keyword_matches: int
) -> Dict:
    """
    Orchestrates the full processing pipeline for a single transcript file.
    """
    print(f"\n--- Starting processing for: {file_name} ---")
    file_content = download_and_extract_text(bucket_name, file_name)
    if not file_content:
        return {
            "file_name": file_name,
            "is_interview": False,
            "reason": "Failed to download or extract text from the file."
        }
    dialogue_content = clean_and_extract_dialogue_segment(file_content)
    if not dialogue_content:
        return {
            "file_name": file_name,
            "is_interview": False,
            "reason": "The 'Transcript' keyword was not found or no content followed it."
        }
    is_interview_flag = is_job_interview(dialogue_content, interview_keywords, min_keyword_matches)
    if not is_interview_flag:
        return {
            "file_name": file_name,
            "is_interview": False,
            "reason": "The file did not meet the criteria for a job interview (not enough keywords)."
        }
    print("File identified as a job interview based on keywords.")
    # Pass person_a_name as the primary_speaker_known_name
    structured_dialogue = parse_interview_dialogue(dialogue_content, person_a_name)
    if not structured_dialogue:
        return {
            "file_name": file_name,
            "is_interview": True, 
            "reason": "Identified as an interview, but failed to parse the dialogue structure.",
            "dialogue": []
        }
    return {
        "file_name": file_name,
        "is_interview": True,
        "dialogue": structured_dialogue
    }


def process_preparation(bucket_name: str, file_name: str, person_a_name: str, job_keywords: List[str], min_keyword_matches: int):
    # --- CONFIGURATION ---
    YOUR_GCS_BUCKET_NAME = bucket_name
    
    
    # IMPORTANT: Set this to the name of one of the primary speakers as it
    # is expected to appear in the transcript text (e.g., the interviewer).
    # For your example "Francisco Ahijado: So um so yeah...", if Francisco Ahijado is
    # the primary speaker you want to identify consistently, set this variable to "Francisco Ahijado".
    # The script will use this exact casing for this speaker in the output.
    # The other speaker's name will be taken from the transcript as it appears.
    YOUR_PERSON_A_NAME = person_a_name # Example: "Francisco Ahijado" or "Gemini-francisco"

    FILES_TO_PROCESS = [file_name]#[
    #     "30 Minute Interview With DevEngine (Jean Massucatto) - 2025/06/02 09:25 GMT-04:00 - Notes by Gemini-francisco.pdf"
    #     # You can add more file names here from the bucket "transcription_poc_uploads_raw"
    # ]

    

    # --- PROCESSING ---
    all_results = []
    for filename_to_process in FILES_TO_PROCESS:
        result = process_transcript_file(
            YOUR_GCS_BUCKET_NAME,
            filename_to_process,
            YOUR_PERSON_A_NAME, # This is passed as the primary_speaker_known_name
            JOB_INTERVIEW_KEYWORDS,
            MIN_KEYWORD_MATCHES
        )
        all_results.append(result)

    # --- PRINT REPORT TO CONSOLE (Optional, can be removed if only file output is needed) ---
    print("\n\n--- CONSOLE PROCESSING REPORT (also saved to file) ---")
    for res_item in all_results:
        print(f"\n########################################")
        print(f"# File: {res_item.get('file_name', 'N/A')}")
        print(f"########################################")
        print(f"Is Interview?: {res_item.get('is_interview', False)}")
        
        if res_item.get('is_interview'):
            print("Extracted Dialogue:")
            dialogue_list = res_item.get('dialogue', [])
            if dialogue_list:
                for turn_count, turn in enumerate(dialogue_list):
                    speaker = turn.get('speaker', 'Unknown Speaker')
                    text = turn.get('text', 'No text found.').replace('\n', ' ').strip()
                    print(f"  Turn {turn_count + 1} - {speaker}: {text}")
            else:
                 print(f"  Reason for no dialogue: {res_item.get('reason', 'Dialogue parsing failed or dialogue was empty.')}")
        else:
            print(f"Reason: {res_item.get('reason', 'No reason provided.')}")
    print("\n--- End of Console Report ---")
    return all_results

@functions_framework.cloud_event
def process_transcription(cloud_event):
    try:
        data = cloud_event.data
    except Exception as e:
        print(f"Error al procesar el evento: {e}")
    
    
    
    #### BLOQUE DEBUG LOCAL #####
    # data = cloud_event.get("data")
    # data = cloud_event.get("data")
    
    # finally:
    #     data = cloud_event["data"]
    
    # YOUR_PERSON_A_NAME = "Francisco Ahijado" # Example: "Francisco Ahijado" or "Gemini-francisco"

    # FILES_TO_PROCESS = "30 Minute Interview With DevEngine (Jean Massucatto) - 2025/06/02 09:25 GMT-04:00 - Notes by Gemini-francisco.pdf"# You can add more file names here from the bucket "transcription_poc_uploads_raw"
    
    bucket_name = BUCKET_NAME#data["bucket"]
    file_name = data["name"]
    candidate_name = extract_candidate_name(file_name)

 
    print(f"Candidate name: {candidate_name}")

    
    print(f"Processing file: gs://{bucket_name}/{file_name}")
    transcript__ = process_preparation(
        bucket_name, file_name, candidate_name, JOB_INTERVIEW_KEYWORDS, MIN_KEYWORD_MATCHES
    )
    
    # tsplit = str(transcript__.splitlines())
    dialogo = transcript__[0].get("dialogue")
    
    if dialogo:
        dialogue_text_str = "\n".join([f"{turn['speaker']}: {turn['text']}" for turn in dialogo])
    else:
        dialogue_text_str = ""
        print("Warning: The dialogue is empty or could not be extracted.")
        # Podrías querer manejar este caso de manera diferente, por ejemplo, no proceder con embeddings o LLM.

    # Generar embeddings del texto (asegúrate de que dialogue_text_str no esté vacío si es un requisito)
    embedding_model = TextEmbeddingModel.from_pretrained(EMBEDDING_MODEL)
    
    # El modelo de embedding espera una lista de textos. Si solo tienes uno, envuélvelo en una lista.
    if dialogue_text_str:
        embeddings = embedding_model.get_embeddings([dialogue_text_str])[0].values
    else:
        embeddings = [] # O maneja el caso de embeddings vacíos como sea apropiado
        print("Warning: No embeddings were generated because the dialogue text is empty.")

    # Guardar el diálogo como .txt en la carpeta "TXT"
    txt_blob_name = f"TXT/{candidate_name}.txt" # Asegúrate que reclut esté definido antes de esta línea
    txt_blob = storage.Client().bucket(BUCKET_DESTINO).blob(txt_blob_name)
    txt_blob.upload_from_string(dialogue_text_str, content_type="text/plain") # Usar dialogue_text_str
    print(f"Dialogue saved at: gs://{BUCKET_DESTINO}/{txt_blob_name}")

    # Guardar los embeddings como .npy en la carpeta "EMBEDDINGS"
    if embeddings: # Solo guardar si se generaron embeddings
        # embeddings_blob_name = f"EMBEDDINGS/{candidate_name}.npy"
        # embeddings_blob = storage.Client().bucket(BUCKET_DESTINO).blob(embeddings_blob_name)
        # embeddings_bytes = BytesIO()
        # np.save(embeddings_bytes, np.array(embeddings))
        # embeddings_bytes.seek(0)
        # embeddings_blob.upload_from_file(embeddings_bytes, content_type="application/octet-stream")
        # print(f"Embeddings guardados en: gs://{BUCKET_DESTINO}/{embeddings_blob_name}")
        # print(f"Embeddings generados para el archivo {file_name}")

        embeddings_summary = embeddings[:40]
        embeddings_str_for_prompt = ", ".join(map(str, embeddings_summary))
    else:
        embeddings_str_for_prompt = "No embeddings generated." # O un string vacío

    with open(PROMPT, "r", encoding="utf-8") as f:
        prompt_template = f.read()
    prompt = prompt_template.format(embeddings_str=embeddings_str_for_prompt) # Usar la variable correcta

    # Llamar a Gemini con el contenido y el prompt
    model = GenerativeModel(LLM_MODEL)
    response = model.generate_content(
        prompt + "\n\nInterview transcription:\n" + dialogue_text_str, # Usar dialogue_text_str
        generation_config={
            "temperature": TEMPERATURE,
            "max_output_tokens": 3500,
            "top_k": TOP_K,
            "top_p": TOP_P,
        }
    )

    # Parsear y guardar o loggear la respuesta
    try:
        candidates = response.to_dict().get("candidates")[0]
        texto = candidates.get("content").get("parts")[0].get("text")
        texto_ = texto.replace("```json\n","").replace("\n","").replace("  ","").replace("```","")
        json_return = json.loads(texto_)
        
    except Exception:
        result_json = {"raw_response": response.text}
    reclut = extract_reclut(file_name)
    output_blob_name = str(candidate_name) + "-" + str(reclut) + ".docx"#.replace(".pdf",".docx")#.replace(".docx", "-interview-analysis.docx")
    document = Document()
    document.add_heading("Interview Analysis", 0)
    if len(list(json_return.keys())) == 1:
        main_key = list(json_return.keys())[0]
        json_iter = json_return.get(main_key)
        for item in json_iter:
            document.add_heading(item, level=1)
            document.add_paragraph(str(json_iter.get(item)))
    else:
        for item in json_return.keys():
            document.add_heading(item, level=1)
            document.add_paragraph(str(json_return.get(item)))

    # Guardar en un buffer
    
    word_buffer = BytesIO()
    document.save(word_buffer)
    word_buffer.seek(0)
    # Save the Word document locally
    # local_output_path = os.path.join(os.getcwd(), output_blob_name)
    # with open(local_output_path, "wb") as local_file:
    #     local_file.write(word_buffer.getvalue())
    # print(f"Resultado guardado localmente en: {local_output_path}")

    # Subir al bucket destino
    storage_client = storage.Client()
    output_blob = storage_client.bucket(BUCKET_DESTINO).blob(output_blob_name)
    output_blob.upload_from_file(word_buffer, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    print(f"Result saved at: gs://{BUCKET_DESTINO}/{output_blob_name}")
    
